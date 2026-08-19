from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import (
    BaseDocTemplate,
    Frame,
    Image,
    KeepTogether,
    PageBreak,
    PageTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = ROOT / "output" / "doc" / "BFC_Instructor_Portal_Introduction_Facilitator_Guide.docx"
PDF_PATH = ROOT / "output" / "pdf" / "BFC_Instructor_Portal_Introduction_Facilitator_Guide.pdf"
ICON_PATH = ROOT / "public" / "pwa-icon-512.png"

NAVY = "174E7A"
BLUE = "1F6EAA"
SKY = "DCEEF8"
PALE = "F3F7FA"
INK = "182634"
MUTED = "526576"
GREEN = "1F7A5A"
AMBER = "A8600A"
RED = "A33A36"
WHITE = "FFFFFF"


AGENDA = [
    ("0-5 min", "Set the purpose", "Why the portal exists, what changes, and the one daily workflow."),
    ("5-10 min", "Sign in and orient", "Home, navigation, phone layout, account and notification preferences."),
    ("10-18 min", "Duty Clock", "Start duty, breaks, reminders, clock-out, history and the booking fallback."),
    ("18-30 min", "Calendar and bookings", "Day view, booking details, warnings, supervision, statuses and notifications."),
    ("30-50 min", "Complete one lesson", "Open the booking, log the flight, allocate time, complete and review the lesson record."),
    ("50-58 min", "Records and logbook", "Pilot file, course progress, outstanding records, links and personal logbook notes."),
    ("58-64 min", "Safety and currency", "Recency, reviews, instructor compliance and what to do with a warning."),
    ("64-68 min", "Notifications", "Bell, push notifications, booking links and notification preferences."),
    ("68-72 min", "Feedback and support", "Capture or upload a screenshot, mark it up and report an issue."),
    ("72-75 min", "Practice and close", "Three-minute participant task, questions, responsibilities and next steps."),
]


DEMO_SECTIONS = [
    {
        "title": "1. Set the purpose",
        "time": "0-5 minutes",
        "purpose": "Give instructors a reason to care before showing screens.",
        "say": (
            "The portal is our shared operational record. It connects the calendar, duty, student training, "
            "flight logging, currency and notifications so we are not relying on separate notes or memory. "
            "Today is about the instructor workflow, not every administrative setting."
        ),
        "show": [
            "Put the portal address on screen: portal.bendigoflyingclub.com.au.",
            "Show the five-step daily flow card on page 2 of this guide.",
            "Explain that portal warnings support judgement and compliance; they do not replace instructor responsibility or current operational requirements.",
        ],
        "ask": "Who has already signed in on both a computer and phone? Note anyone needing account help for the end.",
        "watch": "Do not begin with Settings, reports or billing. Establish the normal instructor day first.",
    },
    {
        "title": "2. Sign in and orient",
        "time": "5-10 minutes",
        "purpose": "Make instructors confident that they can find the daily tools on desktop and phone.",
        "say": (
            "The portal shows features according to your role. On a phone, the most-used destinations sit in the bottom navigation; "
            "the remaining destinations are under More."
        ),
        "show": [
            "Sign in and identify Home, Calendar, Members, Duty, Training Courses, Pilot File, Outstanding Records, My Logbook and Safety.",
            "Open the profile/preferences area and point out password reset, personal details, theme and notification preferences.",
            "On a phone, show the installed PWA or explain Add to Home Screen. Accepting push notifications is optional and can be changed later.",
            "Show searchable dropdown behaviour: typing 'R' narrows to names beginning with R; typing 'Ro' narrows further.",
        ],
        "ask": "Ask one instructor to identify where they would go to find today's bookings and where they would find their own records.",
        "watch": "Keep password or account troubleshooting for the parking lot unless it affects the whole group.",
    },
    {
        "title": "3. Duty Clock - the first operational action",
        "time": "10-18 minutes",
        "purpose": "Make the authoritative duty record and fallback rule unambiguous.",
        "say": (
            "Start the Duty Clock when duty actually starts, record required breaks, and clock out when duty finishes. "
            "Duty time and flight time are different. The Duty Clock is authoritative. If no Duty Clock record exists for the day, "
            "the portal estimates duty from 30 minutes before the first relevant booking until 30 minutes after the last relevant booking."
        ),
        "show": [
            "Open Duty and start a duty period. Explain the fit-for-duty declaration, location and any off-site note.",
            "Show Start break and End break. Explain that the phone reminder arrives 30 minutes before a break is due if no break has been recorded.",
            "Show End duty and the previous-duty cards. Open one past period to show start, finish, breaks and detail.",
            "Point out Sunday-to-Saturday totals: flying hours, total duty time and duty time minus breaks.",
            "Explain that a booking-derived estimate cannot block a real Duty Clock start; the real clock-in replaces the estimate.",
        ],
        "ask": "What is the first thing you should do when you begin duty? Expected answer: start the Duty Clock.",
        "watch": "Do not describe booking duration as duty time. If an estimate is wrong, record or amend the actual duty period.",
    },
    {
        "title": "4. Calendar and bookings",
        "time": "18-30 minutes",
        "purpose": "Teach instructors to read a booking before changing it and to understand warnings.",
        "say": (
            "The calendar is the operational starting point. Open the booking, confirm the student, aircraft, instructor, lesson and status, "
            "then read any warning in context before continuing."
        ),
        "show": [
            "Use Day view first. Identify aircraft rows, times, booking status and the refresh action.",
            "Open a confirmed booking and point out student, instructor, aircraft, description, course/lesson and supervising instructor when applicable.",
            "Create or edit a demonstration booking. Show searchable selectors and the reason field required when continuing past a warning.",
            "Explain the main warning groups: duty/rest, safety/currency, aircraft availability and supervision.",
            "Clarify that solo-hire recency warnings do not apply as solo restrictions when a current instructor is actually on board.",
            "Show that a notification for a booking opens the correct date in Day view and highlights the booking.",
        ],
        "ask": "Before accepting a warning, what should you check? Expected: the facts, the assigned instructor/supervisor, logged records and the reason for continuing.",
        "watch": "Never demonstrate bypassing a warning with a vague reason. Use a genuine, auditable explanation or correct the source data.",
    },
    {
        "title": "5. Complete one lesson end to end",
        "time": "30-50 minutes",
        "purpose": "Demonstrate the highest-value workflow as one continuous task.",
        "say": (
            "After the flight, finish the operational record while the details are fresh: log the flight once, allocate the time correctly, "
            "complete the lesson record, and send it for the required review or acknowledgement."
        ),
        "show": [
            "A. Open the demonstration booking. Confirm that it links to the correct student, course and lesson.",
            "B. Choose Log flight. The portal checks for an existing flight log before allowing another one.",
            "C. Enter actual aircraft and flight details. Accurately record the applicable aircraft meter readings.",
            "D. Allocate dual and solo student time. A flight may contain both. The instructor is PIC for the instructional portion; the student receives dual time. A genuine student solo portion is recorded as solo for the student.",
            "E. Save and show that the calendar status refreshes.",
            "F. Open the linked course lesson and complete the training record: result, competencies, instructor comments and required details.",
            "G. Demonstrate Grammar as a strict grammar-only correction. Demonstrate Rewrite as a clearer, more flexible rewrite that uses student and lesson context without adding fluff or changing the meaning.",
            "H. Create a review draft if required, then complete the review/approval path. Explain the student acknowledgement email and the reduced lesson view after acknowledgement.",
            "I. Briefly explain flight tests: a passed course flight test marks prior course lessons passed, brings course progress to 100 percent and triggers the appropriate completion/pilot-status outcome.",
        ],
        "ask": "Ask an instructor to describe the correct time allocation when they instruct a student: instructor PIC, student dual unless a genuine solo portion is separately recorded.",
        "watch": "Do not invent meter readings, duplicate a flight log or use Rewrite to exaggerate performance. Use a designated demonstration record.",
    },
    {
        "title": "6. Records, outstanding work and logbook",
        "time": "50-58 minutes",
        "purpose": "Show where instructors check completeness after the immediate lesson is finished.",
        "say": "The booking is not the whole record. Use the member profile, outstanding list and logbook to confirm the record is complete and connected.",
        "show": [
            "Open Members, choose the demonstration member and show Courses, Pilot File, Documents, Flight Reviews and Logbook.",
            "Show that completed courses are collapsed by default and cancelled review entries are not presented as current records.",
            "Open Outstanding Records and explain how to find lessons or records awaiting instructor action.",
            "Open My Logbook. Point out date, aircraft, instructor/student names, lesson, booking description and linked destinations.",
            "Explain role allocation: instructor instructional time appears as PIC; the student receives dual. Personal notes can be read by authorised staff but only the logbook owner can add or remove them.",
        ],
        "ask": "Where do you go if you remember completing a flight but are unsure whether the lesson record was finished? Expected: Outstanding Records and the member's course.",
        "watch": "Do not treat a blank personal note as a missing training record; they serve different purposes.",
    },
    {
        "title": "7. Safety, recency and instructor currency",
        "time": "58-64 minutes",
        "purpose": "Explain how operational data feeds safety prompts without overwhelming the meeting.",
        "say": (
            "Safety and currency prompts depend on the records in the portal. If a recent flight is missing, correct the flight record rather than repeatedly acknowledging an inaccurate warning."
        ),
        "show": [
            "Open Safety and show pilot recency/currency information. Archived users are excluded.",
            "Show where an instructor can review their own Pilot File, RAAus BFR date, CASA AFR date, documents and details.",
            "Briefly identify instructor S&P checks and instructor renewal records. Explain that the portal tracks due dates, but current regulatory and club requirements remain authoritative.",
            "Show how an acknowledged safety item is recorded and why the reason must be specific.",
        ],
        "ask": "If the portal says a pilot has not flown recently but you know they flew last week, what should you do first? Expected: find and correct the missing or incorrectly linked flight log.",
        "watch": "Avoid turning this into a full regulatory briefing. Park policy questions that need CFI confirmation.",
    },
    {
        "title": "8. Notifications and phone use",
        "time": "64-68 minutes",
        "purpose": "Make follow-up actions visible without creating notification fatigue.",
        "say": "Notifications should lead to an action. Open the item, deal with it, then mark it read or clear it.",
        "show": [
            "Open the bell and identify the unread number. Mark items read and use Read and clear all when appropriate.",
            "Open a booking notification to show the calendar day and booking highlight.",
            "On the installed PWA, show notification preferences and the company-branded phone notification.",
            "Explain that notification permission is prompted once and can be changed later in portal settings and device settings.",
        ],
        "ask": "Ask everyone to check whether portal notifications are enabled on the device they will actually carry.",
        "watch": "Do not ask people to expose private notification previews on the projector.",
    },
    {
        "title": "9. Feedback and support",
        "time": "68-72 minutes",
        "purpose": "Give instructors a low-friction way to report bugs during rollout.",
        "say": "During rollout, report what happened from the page where it happened. A useful report includes the action, expected result and actual result.",
        "show": [
            "On desktop, select the floating Feedback button in the bottom-left corner.",
            "Choose a tab, window or screen from the browser sharing picker, or use Upload a screenshot if sharing is blocked.",
            "Crop the relevant area, draw on it, select Bug/Improvement/Other, add a clear comment and submit.",
            "Explain that the report includes page address, browser and screen size and is sent to Lincoln.",
        ],
        "ask": "What makes a useful bug report? Expected: what you were doing, what you expected, what happened, and the affected booking/member if relevant.",
        "watch": "Capture only what is needed. Avoid unrelated personal or sensitive information in screenshots.",
    },
    {
        "title": "10. Practice, questions and close",
        "time": "72-75 minutes",
        "purpose": "Finish with recall and responsibility, not another feature tour.",
        "say": (
            "The standard is simple: start duty, read the booking, log the flight once, finish the training record, act on warnings and outstanding items, then clock out."
        ),
        "show": [
            "Ask participants to complete the three-minute practice on page 10.",
            "Review the daily quick reference on the final page.",
            "Confirm the support route and when instructors are expected to begin using the system.",
        ],
        "ask": "What is the one part of the workflow you want to practise again before go-live? Record answers in the parking lot.",
        "watch": "Do not let the meeting end without confirming ownership of duty, flight logging and training-record completion.",
    },
]


PRE_MEETING = [
    "Confirm the meeting date, room, screen and 75-minute allocation.",
    "Confirm stable internet and have a phone hotspot available as backup.",
    "Sign in as a presenter before the meeting and verify the portal opens on desktop and phone.",
    "Use a designated demonstration member or your own test profile. Do not alter a real student's operational record for the demonstration.",
    "Prepare one future confirmed booking linked to a demonstration student, aircraft, course and lesson.",
    "Ensure the demonstration booking does not already have a flight log unless duplicate prevention is the point being shown.",
    "Prepare safe example aircraft times/meter values and a short lesson comment for Grammar and Rewrite.",
    "Open these tabs in advance: Home, Duty, Calendar Day view, demonstration member profile and Training Courses.",
    "Install/open the PWA on your phone and check that notification permission is enabled.",
    "Close unrelated tabs, email, messages and sensitive member information before screen sharing.",
    "Set browser zoom to 100 percent and silence unrelated phone/desktop notifications.",
    "Print one copy of this guide and have the PDF available offline.",
]


PRACTICE = [
    "Open the Calendar in Day view and locate a booking.",
    "Open the booking and identify the student, aircraft, instructor, lesson and status.",
    "Navigate to the member's course and locate the linked lesson.",
    "Open Duty and identify where to start duty, record a break and view a prior period.",
    "Open My Logbook and follow one link from an entry.",
    "Open Feedback, identify both capture methods, then close it without submitting unless a meeting test is intended.",
]


DAILY_FLOW = [
    ("1", "Start duty", "Use the Duty Clock at the actual start of duty. Record required breaks."),
    ("2", "Read the booking", "Confirm people, aircraft, lesson, supervision, status and warnings."),
    ("3", "Log the flight once", "Enter actual details and allocate PIC, dual and solo correctly."),
    ("4", "Finish the lesson record", "Record performance clearly; complete review and acknowledgement steps."),
    ("5", "Close the day", "Resolve outstanding items, check notifications and clock out at actual finish."),
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    run.font.size = Pt(8)
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr_text)
    run._r.append(fld_char_2)


def add_docx_label(document: Document, label: str, text: str, colour: str) -> None:
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(2.1)
    table.columns[1].width = Cm(15.1)
    left, right = table.rows[0].cells
    set_cell_shading(left, colour)
    set_cell_shading(right, PALE)
    for cell in (left, right):
        set_cell_margins(cell, 100, 150, 100, 150)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = left.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(label.upper())
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(255, 255, 255)
    p = right.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor.from_string(INK)
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_docx_bullets(document: Document, items: list[str], checklist: bool = False) -> None:
    for item in items:
        p = document.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Cm(0.55)
        p.paragraph_format.first_line_indent = Cm(-0.25)
        p.paragraph_format.space_after = Pt(4)
        text = f"[ ] {item}" if checklist else item
        run = p.add_run(text)
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor.from_string(INK)


def add_docx_daily_flow(document: Document) -> None:
    table = document.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [Cm(1.2), Cm(4.3), Cm(11.7)]
    for idx, width in enumerate(widths):
        table.columns[idx].width = width
    header = table.rows[0].cells
    for i, label in enumerate(("Step", "Action", "Standard")):
        set_cell_shading(header[i], NAVY)
        set_cell_margins(header[i])
        run = header[i].paragraphs[0].add_run(label)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(9)
    set_repeat_table_header(table.rows[0])
    for number, action, standard in DAILY_FLOW:
        cells = table.add_row().cells
        values = (number, action, standard)
        for i, value in enumerate(values):
            set_cell_margins(cells[i])
            set_cell_shading(cells[i], WHITE if int(number) % 2 else PALE)
            run = cells[i].paragraphs[0].add_run(value)
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor.from_string(INK)
            if i < 2:
                run.bold = True


def build_docx() -> None:
    DOCX_PATH.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    section = document.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21)
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(9.5)
    styles["Normal"].font.color.rgb = RGBColor.from_string(INK)
    for name, size, colour in (("Title", 28, NAVY), ("Heading 1", 18, NAVY), ("Heading 2", 13, BLUE)):
        style = styles[name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(colour)
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(7)

    if ICON_PATH.exists():
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(ICON_PATH), width=Inches(1.25))
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("BENDIGO FLYING CLUB")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor.from_string(BLUE)
    p = document.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Instructor Portal Introduction")
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("FACILITATOR GUIDE AND RUN SHEET")
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor.from_string(MUTED)
    document.add_paragraph()

    cover = document.add_table(rows=4, cols=2)
    cover.alignment = WD_TABLE_ALIGNMENT.CENTER
    labels = ("Meeting", "Recommended duration", "Presenter", "Room / attendees")
    values = ("Monday, 17 August 2026 at 1:00 pm", "75 minutes - 60 minute core plus 15 minutes practice and questions", "Lincoln Cottingham", "_______________________________________________")
    for row, label, value in zip(cover.rows, labels, values):
        set_cell_shading(row.cells[0], SKY)
        set_cell_shading(row.cells[1], PALE)
        for cell in row.cells:
            set_cell_margins(cell, 140, 180, 140, 180)
        r = row.cells[0].paragraphs[0].add_run(label)
        r.bold = True
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor.from_string(NAVY)
        r = row.cells[1].paragraphs[0].add_run(value)
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor.from_string(INK)

    document.add_paragraph()
    add_docx_label(
        document,
        "Outcome",
        "Every instructor should leave able to start duty, read a booking, complete a lesson record, find outstanding work and report a problem.",
        GREEN,
    )
    add_docx_label(
        document,
        "Key message",
        "Start duty. Read the booking. Log the flight once. Finish the training record. Act on warnings. Clock out.",
        NAVY,
    )
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(22)
    r = p.add_run("Presenter copy - use the SAY / SHOW / ASK / WATCH prompts during the live demonstration")
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string(MUTED)

    footer = section.footer.paragraphs[0]
    add_page_number(footer)

    document.add_page_break()
    document.add_heading("The order to show instructors", level=1)
    p = document.add_paragraph("Lead with the instructor's working day. Do not follow the sidebar from top to bottom.")
    p.paragraph_format.space_after = Pt(8)
    p.runs[0].font.size = Pt(10)
    p.runs[0].font.color.rgb = RGBColor.from_string(MUTED)
    add_docx_daily_flow(document)
    document.add_paragraph()
    add_docx_label(document, "Rule", "If time runs short, complete Duty Clock, Calendar and the end-to-end lesson. Those are the non-negotiable demonstrations.", AMBER)

    document.add_heading("Timed run sheet", level=2)
    table = document.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [Cm(2.3), Cm(4.3), Cm(10.6)]
    for idx, width in enumerate(widths):
        table.columns[idx].width = width
    for i, label in enumerate(("Time", "Show", "Instructor outcome")):
        set_cell_shading(table.rows[0].cells[i], NAVY)
        set_cell_margins(table.rows[0].cells[i])
        r = table.rows[0].cells[i].paragraphs[0].add_run(label)
        r.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
        r.font.size = Pt(8.5)
    set_repeat_table_header(table.rows[0])
    for index, row_data in enumerate(AGENDA):
        cells = table.add_row().cells
        for i, value in enumerate(row_data):
            set_cell_margins(cells[i], 80, 110, 80, 110)
            set_cell_shading(cells[i], WHITE if index % 2 == 0 else PALE)
            r = cells[i].paragraphs[0].add_run(value)
            r.font.size = Pt(8.2)
            r.font.color.rgb = RGBColor.from_string(INK)
            if i < 2:
                r.bold = True

    document.add_page_break()
    document.add_heading("Before the meeting", level=1)
    add_docx_label(document, "Aim", "Arrive with a safe demonstration record and no setup tasks competing with teaching time.", BLUE)
    add_docx_bullets(document, PRE_MEETING, checklist=True)
    document.add_heading("Prepare the demonstration story", level=2)
    story = [
        "A demonstration instructor begins duty.",
        "They open a confirmed lesson booking and check the warning state.",
        "After the flight they log it once and allocate instructor PIC and student dual/solo correctly.",
        "They complete the linked lesson record, use Grammar and Rewrite appropriately, and send it through review/acknowledgement.",
        "They confirm the result in Outstanding Records and the logbook, then clock out.",
    ]
    add_docx_bullets(document, story)
    add_docx_label(document, "Privacy", "Use a demonstration profile where possible. If a real record must be shown, minimise personal information and do not alter it for demonstration purposes.", RED)

    for section_data in DEMO_SECTIONS:
        document.add_page_break()
        document.add_heading(section_data["title"], level=1)
        p = document.add_paragraph()
        r = p.add_run(section_data["time"])
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor.from_string(BLUE)
        p.add_run(f"  |  {section_data['purpose']}")
        add_docx_label(document, "Say", section_data["say"], NAVY)
        document.add_heading("Show", level=2)
        add_docx_bullets(document, section_data["show"])
        add_docx_label(document, "Ask", section_data["ask"], GREEN)
        add_docx_label(document, "Watch", section_data["watch"], AMBER)

    document.add_page_break()
    document.add_heading("Three-minute participant practice", level=1)
    p = document.add_paragraph(
        "Ask each instructor to use their own account. Pair anyone who is not yet activated with an instructor who is signed in. "
        "Use demonstration data and do not save a false operational record."
    )
    p.paragraph_format.space_after = Pt(8)
    add_docx_bullets(document, PRACTICE, checklist=True)
    document.add_heading("Questions and parking lot", level=2)
    parking = document.add_table(rows=6, cols=3)
    parking.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = [Cm(7.7), Cm(5.5), Cm(4.0)]
    for idx, width in enumerate(widths):
        parking.columns[idx].width = width
    for i, value in enumerate(("Question / issue", "Owner / answer", "Follow-up date")):
        set_cell_shading(parking.rows[0].cells[i], NAVY)
        r = parking.rows[0].cells[i].paragraphs[0].add_run(value)
        r.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
        r.font.size = Pt(9)
    for row in parking.rows:
        for cell in row.cells:
            set_cell_margins(cell, 120, 120, 160, 120)

    document.add_page_break()
    document.add_heading("Daily instructor quick reference", level=1)
    add_docx_daily_flow(document)
    document.add_heading("Before the first flight", level=2)
    add_docx_bullets(document, [
        "Start the Duty Clock at actual duty commencement.",
        "Review the calendar, aircraft, student, lesson, supervision and warnings.",
        "Resolve missing or incorrect source records instead of repeatedly acknowledging a false warning.",
    ], checklist=True)
    document.add_heading("After every flight", level=2)
    add_docx_bullets(document, [
        "Log the flight once and record actual aircraft/meter details.",
        "Allocate instructor PIC, student dual and any genuine student solo portion correctly.",
        "Complete the linked training record before details are forgotten.",
        "Use Grammar for grammar only; use Rewrite for a clear, faithful rewrite without fluff.",
        "Complete required review, approval and acknowledgement steps.",
    ], checklist=True)
    document.add_heading("Before leaving", level=2)
    add_docx_bullets(document, [
        "Check Outstanding Records and notifications.",
        "Record/end any required break and clock out at actual duty finish.",
        "Use Feedback for a bug or improvement and include a useful screenshot and explanation.",
    ], checklist=True)
    add_docx_label(document, "Fallback", "If no Duty Clock record exists, the portal estimates duty from 30 minutes before the first relevant booking to 30 minutes after the last. Record actual duty whenever possible.", AMBER)
    add_docx_label(document, "Support", "Portal feedback: bottom-left Feedback button on desktop. Direct contact: lincoln@bbkm.com.au", BLUE)

    document.save(DOCX_PATH)


def pdf_styles():
    sample = getSampleStyleSheet()
    return {
        "title": ParagraphStyle("TitleBFC", parent=sample["Title"], fontName="Helvetica-Bold", fontSize=24, leading=29, textColor=colors.HexColor(f"#{NAVY}"), alignment=TA_CENTER, spaceAfter=8),
        "subtitle": ParagraphStyle("SubtitleBFC", parent=sample["Normal"], fontName="Helvetica-Bold", fontSize=10, leading=13, textColor=colors.HexColor(f"#{MUTED}"), alignment=TA_CENTER, spaceAfter=10),
        "h1": ParagraphStyle("H1BFC", parent=sample["Heading1"], fontName="Helvetica-Bold", fontSize=17, leading=21, textColor=colors.HexColor(f"#{NAVY}"), spaceAfter=8),
        "h2": ParagraphStyle("H2BFC", parent=sample["Heading2"], fontName="Helvetica-Bold", fontSize=11.5, leading=14, textColor=colors.HexColor(f"#{BLUE}"), spaceBefore=7, spaceAfter=5),
        "body": ParagraphStyle("BodyBFC", parent=sample["BodyText"], fontName="Helvetica", fontSize=8.7, leading=12, textColor=colors.HexColor(f"#{INK}"), spaceAfter=5),
        "small": ParagraphStyle("SmallBFC", parent=sample["BodyText"], fontName="Helvetica", fontSize=7.5, leading=10, textColor=colors.HexColor(f"#{MUTED}")),
        "bullet": ParagraphStyle("BulletBFC", parent=sample["BodyText"], fontName="Helvetica", fontSize=8.4, leading=11.2, leftIndent=11, firstLineIndent=-7, textColor=colors.HexColor(f"#{INK}"), spaceAfter=3),
        "label": ParagraphStyle("LabelBFC", parent=sample["BodyText"], fontName="Helvetica-Bold", fontSize=7.5, leading=9, textColor=colors.white, alignment=TA_CENTER),
        "callout": ParagraphStyle("CalloutBFC", parent=sample["BodyText"], fontName="Helvetica", fontSize=8.3, leading=11, textColor=colors.HexColor(f"#{INK}")),
    }


def pdf_callout(label: str, text: str, colour: str, styles) -> Table:
    table = Table(
        [[Paragraph(label.upper(), styles["label"]), Paragraph(text, styles["callout"])]],
        colWidths=[22 * mm, 150 * mm],
    )
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (0, 0), colors.HexColor(f"#{colour}")),
        ("BACKGROUND", (1, 0), (1, 0), colors.HexColor(f"#{PALE}")),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("BOX", (0, 0), (-1, -1), 0.4, colors.HexColor(f"#{SKY}")),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
    ]))
    return table


def pdf_bullets(items: list[str], styles, checklist: bool = False):
    prefix = "[ ]" if checklist else "-"
    return [Paragraph(f"{prefix} {item}", styles["bullet"]) for item in items]


def pdf_daily_flow(styles) -> Table:
    data = [[Paragraph("Step", styles["small"]), Paragraph("Action", styles["small"]), Paragraph("Standard", styles["small"])]]
    for number, action, standard in DAILY_FLOW:
        data.append([
            Paragraph(f"<b>{number}</b>", styles["body"]),
            Paragraph(f"<b>{action}</b>", styles["body"]),
            Paragraph(standard, styles["body"]),
        ])
    table = Table(data, colWidths=[13 * mm, 43 * mm, 116 * mm], repeatRows=1)
    style = [
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(f"#{NAVY}")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor(f"#{SKY}")),
        ("LEFTPADDING", (0, 0), (-1, -1), 5),
        ("RIGHTPADDING", (0, 0), (-1, -1), 5),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
    ]
    for row in range(1, len(data)):
        style.append(("BACKGROUND", (0, row), (-1, row), colors.white if row % 2 else colors.HexColor(f"#{PALE}")))
    table.setStyle(TableStyle(style))
    return table


class NumberedCanvasMixin:
    pass


def build_pdf() -> None:
    PDF_PATH.parent.mkdir(parents=True, exist_ok=True)
    styles = pdf_styles()
    doc = BaseDocTemplate(
        str(PDF_PATH),
        pagesize=A4,
        leftMargin=19 * mm,
        rightMargin=19 * mm,
        topMargin=18 * mm,
        bottomMargin=16 * mm,
        title="BFC Instructor Portal Introduction Facilitator Guide",
        author="Bendigo Flying Club",
    )

    def draw_page(canvas, document):
        canvas.saveState()
        if document.page > 1:
            canvas.setStrokeColor(colors.HexColor(f"#{SKY}"))
            canvas.line(19 * mm, 282 * mm, 191 * mm, 282 * mm)
            canvas.setFont("Helvetica-Bold", 7.2)
            canvas.setFillColor(colors.HexColor(f"#{NAVY}"))
            canvas.drawString(19 * mm, 286 * mm, "BENDIGO FLYING CLUB - INSTRUCTOR PORTAL INTRODUCTION")
        canvas.setFont("Helvetica", 7)
        canvas.setFillColor(colors.HexColor(f"#{MUTED}"))
        canvas.drawRightString(191 * mm, 9 * mm, f"Page {document.page}")
        canvas.restoreState()

    frame = Frame(doc.leftMargin, doc.bottomMargin, doc.width, doc.height, id="normal")
    doc.addPageTemplates(PageTemplate(id="all", frames=[frame], onPage=draw_page))
    story = []

    if ICON_PATH.exists():
        logo = Image(str(ICON_PATH), width=27 * mm, height=27 * mm)
        logo.hAlign = "CENTER"
        story.extend([Spacer(1, 8 * mm), logo, Spacer(1, 3 * mm)])
    story.append(Paragraph("BENDIGO FLYING CLUB", styles["subtitle"]))
    story.append(Paragraph("Instructor Portal Introduction", styles["title"]))
    story.append(Paragraph("FACILITATOR GUIDE AND RUN SHEET", styles["subtitle"]))
    story.append(Spacer(1, 7 * mm))
    cover_data = [
        ["Meeting", "Monday, 17 August 2026 at 1:00 pm"],
        ["Recommended duration", "75 minutes - 60 minute core plus 15 minutes practice and questions"],
        ["Presenter", "Lincoln Cottingham"],
        ["Room / attendees", "_______________________________________________"],
    ]
    cover_table = Table(
        [[Paragraph(f"<b>{label}</b>", styles["body"]), Paragraph(value, styles["body"])] for label, value in cover_data],
        colWidths=[43 * mm, 119 * mm],
    )
    cover_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (0, -1), colors.HexColor(f"#{SKY}")),
        ("BACKGROUND", (1, 0), (1, -1), colors.HexColor(f"#{PALE}")),
        ("GRID", (0, 0), (-1, -1), 0.4, colors.white),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, -1), 7),
        ("RIGHTPADDING", (0, 0), (-1, -1), 7),
        ("TOPPADDING", (0, 0), (-1, -1), 7),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
    ]))
    story.extend([cover_table, Spacer(1, 7 * mm)])
    story.append(pdf_callout("Outcome", "Every instructor should leave able to start duty, read a booking, complete a lesson record, find outstanding work and report a problem.", GREEN, styles))
    story.append(Spacer(1, 3 * mm))
    story.append(pdf_callout("Key message", "Start duty. Read the booking. Log the flight once. Finish the training record. Act on warnings. Clock out.", NAVY, styles))
    story.append(Spacer(1, 12 * mm))
    story.append(Paragraph("Presenter copy - use the SAY / SHOW / ASK / WATCH prompts during the live demonstration", styles["subtitle"]))

    story.append(PageBreak())
    story.append(Paragraph("The order to show instructors", styles["h1"]))
    story.append(Paragraph("Lead with the instructor's working day. Do not follow the sidebar from top to bottom.", styles["body"]))
    story.append(pdf_daily_flow(styles))
    story.append(Spacer(1, 3 * mm))
    story.append(pdf_callout("Rule", "If time runs short, complete Duty Clock, Calendar and the end-to-end lesson. Those are the non-negotiable demonstrations.", AMBER, styles))
    story.append(Paragraph("Timed run sheet", styles["h2"]))
    agenda_data = [[Paragraph("Time", styles["small"]), Paragraph("Show", styles["small"]), Paragraph("Instructor outcome", styles["small"])]]
    for time, title, outcome in AGENDA:
        agenda_data.append([Paragraph(f"<b>{time}</b>", styles["small"]), Paragraph(f"<b>{title}</b>", styles["small"]), Paragraph(outcome, styles["small"])])
    agenda_table = Table(agenda_data, colWidths=[22 * mm, 40 * mm, 110 * mm], repeatRows=1)
    agenda_style = [
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(f"#{NAVY}")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("GRID", (0, 0), (-1, -1), 0.3, colors.HexColor(f"#{SKY}")),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ("TOPPADDING", (0, 0), (-1, -1), 3.5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3.5),
    ]
    for row in range(1, len(agenda_data)):
        agenda_style.append(("BACKGROUND", (0, row), (-1, row), colors.white if row % 2 else colors.HexColor(f"#{PALE}")))
    agenda_table.setStyle(TableStyle(agenda_style))
    story.append(agenda_table)

    story.append(PageBreak())
    story.append(Paragraph("Before the meeting", styles["h1"]))
    story.append(pdf_callout("Aim", "Arrive with a safe demonstration record and no setup tasks competing with teaching time.", BLUE, styles))
    story.append(Spacer(1, 2 * mm))
    story.extend(pdf_bullets(PRE_MEETING, styles, checklist=True))
    story.append(Paragraph("Prepare the demonstration story", styles["h2"]))
    story.extend(pdf_bullets([
        "A demonstration instructor begins duty.",
        "They open a confirmed lesson booking and check the warning state.",
        "After the flight they log it once and allocate instructor PIC and student dual/solo correctly.",
        "They complete the lesson record, use Grammar and Rewrite appropriately, and send it through review/acknowledgement.",
        "They confirm the result in Outstanding Records and the logbook, then clock out.",
    ], styles))
    story.append(pdf_callout("Privacy", "Use a demonstration profile where possible. If a real record must be shown, minimise personal information and do not alter it for demonstration purposes.", RED, styles))

    for section_data in DEMO_SECTIONS:
        story.append(PageBreak())
        story.append(Paragraph(section_data["title"], styles["h1"]))
        story.append(Paragraph(f"<b>{section_data['time']}</b>  |  {section_data['purpose']}", styles["body"]))
        story.append(pdf_callout("Say", section_data["say"], NAVY, styles))
        story.append(Paragraph("Show", styles["h2"]))
        story.extend(pdf_bullets(section_data["show"], styles))
        story.append(Spacer(1, 2 * mm))
        story.append(pdf_callout("Ask", section_data["ask"], GREEN, styles))
        story.append(Spacer(1, 2 * mm))
        story.append(pdf_callout("Watch", section_data["watch"], AMBER, styles))

    story.append(PageBreak())
    story.append(Paragraph("Three-minute participant practice", styles["h1"]))
    story.append(Paragraph("Ask each instructor to use their own account. Pair anyone who is not yet activated with an instructor who is signed in. Use demonstration data and do not save a false operational record.", styles["body"]))
    story.extend(pdf_bullets(PRACTICE, styles, checklist=True))
    story.append(Paragraph("Questions and parking lot", styles["h2"]))
    parking_data = [[Paragraph("Question / issue", styles["small"]), Paragraph("Owner / answer", styles["small"]), Paragraph("Follow-up date", styles["small"])]]
    parking_data.extend([["", "", ""] for _ in range(5)])
    parking = Table(parking_data, colWidths=[78 * mm, 57 * mm, 37 * mm], rowHeights=[8 * mm] + [15 * mm] * 5)
    parking.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(f"#{NAVY}")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor(f"#{SKY}")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 5),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
    ]))
    story.append(parking)

    story.append(PageBreak())
    story.append(Paragraph("Daily instructor quick reference", styles["h1"]))
    story.append(pdf_daily_flow(styles))
    for heading, items in (
        ("Before the first flight", [
            "Start the Duty Clock at actual duty commencement.",
            "Review the calendar, aircraft, student, lesson, supervision and warnings.",
            "Resolve missing or incorrect source records instead of repeatedly acknowledging a false warning.",
        ]),
        ("After every flight", [
            "Log the flight once and record actual aircraft/meter details.",
            "Allocate instructor PIC, student dual and any genuine student solo portion correctly.",
            "Complete the linked training record before details are forgotten.",
            "Use Grammar for grammar only; use Rewrite for a clear, faithful rewrite without fluff.",
            "Complete required review, approval and acknowledgement steps.",
        ]),
        ("Before leaving", [
            "Check Outstanding Records and notifications.",
            "Record/end any required break and clock out at actual duty finish.",
            "Use Feedback for a bug or improvement and include a useful screenshot and explanation.",
        ]),
    ):
        story.append(Paragraph(heading, styles["h2"]))
        story.extend(pdf_bullets(items, styles, checklist=True))
    story.append(Spacer(1, 2 * mm))
    story.append(pdf_callout("Fallback", "If no Duty Clock record exists, the portal estimates duty from 30 minutes before the first relevant booking to 30 minutes after the last. Record actual duty whenever possible.", AMBER, styles))
    story.append(Spacer(1, 2 * mm))
    story.append(pdf_callout("Support", "Portal feedback: bottom-left Feedback button on desktop. Direct contact: lincoln@bbkm.com.au", BLUE, styles))

    doc.build(story)


if __name__ == "__main__":
    build_docx()
    build_pdf()
    print(DOCX_PATH)
    print(PDF_PATH)
